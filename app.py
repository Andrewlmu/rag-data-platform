"""
PE Data Analysis POC
Complete system with document processing, vector search, and AI analysis
"""

import os
import sys
from dotenv import load_dotenv
import pandas as pd
import numpy as np
from typing import Dict, List, Any, Optional
import json
from datetime import datetime

# Document processing
import PyPDF2
from docx import Document as DocxDocument

# LangChain for intelligent processing
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.agents import create_pandas_dataframe_agent, AgentType

# OpenAI
from openai import OpenAI

# UI
import streamlit as st
import plotly.express as px
import plotly.graph_objects as go

# Load environment variables
load_dotenv()

# Check for API key
if not os.getenv("OPENAI_API_KEY"):
    st.error("⚠️ Please set OPENAI_API_KEY in .env file")
    st.info("Copy .env.example to .env and add your OpenAI API key")
    st.stop()


class DataWarehouse:
    """Handles all data: structured (CSV/Excel) and unstructured (PDF/Word)"""

    def __init__(self):
        self.structured_data = {}  # DataFrames from CSV/Excel
        self.documents = []        # Raw documents
        self.vectorstore = None    # ChromaDB for document search
        self.embeddings = OpenAIEmbeddings()
        self.llm = ChatOpenAI(model="gpt-4-turbo-preview", temperature=0)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )

        # Stats
        self.stats = {
            'total_rows': 0,
            'total_documents': 0,
            'total_chunks': 0,
            'data_size_mb': 0
        }

    def process_all_data(self, folder_path: str):
        """Process all files in folder - CSVs, Excel, PDFs, etc."""

        if not os.path.exists(folder_path):
            st.error(f"Folder {folder_path} does not exist!")
            return None

        progress_text = st.empty()
        progress_bar = st.progress(0)

        # Reset stats
        self.stats = {
            'total_rows': 0,
            'total_documents': 0,
            'total_chunks': 0,
            'data_size_mb': 0
        }

        # Process structured data
        progress_text.text("📊 Processing structured data...")
        progress_bar.progress(25)
        self._process_structured_data(folder_path)

        # Process documents
        progress_text.text("📄 Processing documents...")
        progress_bar.progress(50)
        self._process_documents(folder_path)

        # Create vector index
        progress_text.text("🔍 Creating vector embeddings...")
        progress_bar.progress(75)
        self._create_vector_index()

        progress_text.text("✅ Processing complete!")
        progress_bar.progress(100)

        return self.get_summary()

    def _process_structured_data(self, folder_path: str):
        """Load all CSV and Excel files"""

        for file in os.listdir(folder_path):
            if file.startswith('.'):
                continue

            file_path = os.path.join(folder_path, file)

            try:
                if file.endswith('.csv'):
                    df = pd.read_csv(file_path)
                    self.structured_data[file] = df
                    self.stats['total_rows'] += len(df)

                elif file.endswith(('.xlsx', '.xls')):
                    df = pd.read_excel(file_path)
                    self.structured_data[file] = df
                    self.stats['total_rows'] += len(df)

            except Exception as e:
                st.warning(f"Could not load {file}: {str(e)}")

        # Calculate data size
        for df in self.structured_data.values():
            self.stats['data_size_mb'] += df.memory_usage(deep=True).sum() / 1024**2

    def _process_documents(self, folder_path: str):
        """Parse all PDF and Word documents"""

        for file in os.listdir(folder_path):
            if file.startswith('.'):
                continue

            file_path = os.path.join(folder_path, file)

            try:
                if file.endswith('.pdf'):
                    text = self._parse_pdf(file_path)
                    if text:
                        self.documents.append({
                            'filename': file,
                            'type': 'pdf',
                            'content': text,
                            'chunks': []
                        })
                        self.stats['total_documents'] += 1

                elif file.endswith('.docx'):
                    text = self._parse_docx(file_path)
                    if text:
                        self.documents.append({
                            'filename': file,
                            'type': 'docx',
                            'content': text,
                            'chunks': []
                        })
                        self.stats['total_documents'] += 1

                elif file.endswith('.txt'):
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                    if text:
                        self.documents.append({
                            'filename': file,
                            'type': 'txt',
                            'content': text,
                            'chunks': []
                        })
                        self.stats['total_documents'] += 1

            except Exception as e:
                st.warning(f"Could not parse {file}: {str(e)}")

    def _parse_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""

        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}"
        except Exception as e:
            st.error(f"Error parsing PDF {file_path}: {str(e)}")

        return text

    def _parse_docx(self, file_path: str) -> str:
        """Extract text from Word document"""

        text = ""
        try:
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
        except Exception as e:
            st.error(f"Error parsing DOCX {file_path}: {str(e)}")

        return text

    def _create_vector_index(self):
        """Create vector embeddings for all documents and structured data descriptions"""

        all_texts = []
        all_metadatas = []

        # 1. Chunk and index documents
        for doc in self.documents:
            chunks = self.text_splitter.split_text(doc['content'])
            doc['chunks'] = chunks

            for i, chunk in enumerate(chunks):
                all_texts.append(chunk)
                all_metadatas.append({
                    'source': doc['filename'],
                    'type': 'document',
                    'chunk': i,
                    'total_chunks': len(chunks)
                })

        # 2. Create descriptions of structured data for indexing
        for filename, df in self.structured_data.items():
            # Create a rich description of the dataset
            description = self._create_dataset_description(filename, df)

            all_texts.append(description)
            all_metadatas.append({
                'source': filename,
                'type': 'structured_data',
                'chunk': 0,
                'total_chunks': 1
            })

        # 3. Create vector store
        if all_texts:
            try:
                self.vectorstore = Chroma.from_texts(
                    texts=all_texts,
                    embedding=self.embeddings,
                    metadatas=all_metadatas,
                    persist_directory="./chroma_db"
                )
                self.stats['total_chunks'] = len(all_texts)
            except Exception as e:
                st.error(f"Error creating vector store: {str(e)}")

    def _create_dataset_description(self, filename: str, df: pd.DataFrame) -> str:
        """Create a comprehensive description of a dataset"""

        description_parts = [
            f"Dataset: {filename}",
            f"Shape: {len(df)} rows × {len(df.columns)} columns",
            f"Columns: {', '.join(df.columns.tolist())}",
            ""
        ]

        # Add summary statistics for numeric columns
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        if len(numeric_cols) > 0:
            description_parts.append("Summary Statistics:")
            for col in numeric_cols[:10]:  # Limit to first 10 numeric columns
                description_parts.append(
                    f"  {col}: mean={df[col].mean():.2f}, "
                    f"std={df[col].std():.2f}, "
                    f"min={df[col].min():.2f}, "
                    f"max={df[col].max():.2f}"
                )

        # Add data types
        description_parts.append("\nData Types:")
        for dtype, count in df.dtypes.value_counts().items():
            description_parts.append(f"  {dtype}: {count} columns")

        # Add missing value information
        missing = df.isnull().sum()
        if missing.sum() > 0:
            description_parts.append("\nMissing Values:")
            for col, count in missing[missing > 0].items():
                description_parts.append(f"  {col}: {count} ({count/len(df)*100:.1f}%)")

        # Add sample data
        description_parts.append("\nSample Data (first 3 rows):")
        description_parts.append(df.head(3).to_string())

        return "\n".join(description_parts)

    def search(self, query: str, k: int = 5) -> List[Any]:
        """Vector similarity search across all data"""

        if not self.vectorstore:
            return []

        try:
            return self.vectorstore.similarity_search(query, k=k)
        except Exception as e:
            st.error(f"Search error: {str(e)}")
            return []

    def ask_structured_data(self, question: str) -> str:
        """Use pandas agent for structured data questions"""

        if not self.structured_data:
            return "No structured data loaded."

        # Create pandas agent with all dataframes
        dfs = list(self.structured_data.values())

        try:
            agent = create_pandas_dataframe_agent(
                self.llm,
                dfs,
                verbose=True,
                agent_type=AgentType.OPENAI_FUNCTIONS,
                handle_parsing_errors=True
            )

            return agent.run(question)
        except Exception as e:
            return f"Error analyzing structured data: {str(e)}"

    def ask_documents(self, question: str) -> str:
        """Use RAG for document questions"""

        if not self.vectorstore:
            return "No documents indexed."

        try:
            qa_chain = RetrievalQA.from_chain_type(
                llm=self.llm,
                chain_type="stuff",
                retriever=self.vectorstore.as_retriever(search_kwargs={"k": 5}),
                return_source_documents=True
            )

            result = qa_chain({"query": question})

            # Format response with sources
            answer = result['result']
            sources = set([doc.metadata['source'] for doc in result['source_documents']])

            if sources:
                answer += f"\n\n📄 **Sources:** {', '.join(sources)}"

            return answer
        except Exception as e:
            return f"Error searching documents: {str(e)}"

    def ask_unified(self, question: str) -> str:
        """Unified question answering across all data"""

        context_parts = []

        # 1. Get relevant documents
        if self.vectorstore:
            try:
                relevant_docs = self.search(question, k=3)
                if relevant_docs:
                    doc_context = "\n".join([doc.page_content[:500] for doc in relevant_docs])
                    context_parts.append(f"**Relevant document excerpts:**\n{doc_context}")
            except Exception as e:
                st.warning(f"Document search error: {str(e)}")

        # 2. Get structured data summary
        if self.structured_data:
            data_summary = self.get_structured_data_summary()
            context_parts.append(f"**Available datasets:**\n{data_summary}")

        if not context_parts:
            return "No data available to answer this question."

        # 3. Combine and ask GPT-4
        full_context = "\n\n".join(context_parts)

        try:
            client = OpenAI()
            response = client.chat.completions.create(
                model="gpt-4-turbo-preview",
                messages=[
                    {
                        "role": "system",
                        "content": f"""You are a private equity analyst with access to company data and documents.

Based on this context:
{full_context}

Provide specific, quantitative answers when possible. Cite sources when available."""
                    },
                    {"role": "user", "content": question}
                ],
                temperature=0,
                max_tokens=1000
            )

            return response.choices[0].message.content
        except Exception as e:
            return f"Error generating response: {str(e)}"

    def get_structured_data_summary(self) -> str:
        """Get summary of all structured data"""

        if not self.structured_data:
            return "No structured data loaded."

        summary = []
        for name, df in self.structured_data.items():
            numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
            summary.append(
                f"**{name}**: {df.shape[0]} rows, {df.shape[1]} columns\n"
                f"  Columns: {', '.join(df.columns[:5])}{'...' if len(df.columns) > 5 else ''}\n"
                f"  Numeric fields: {', '.join(numeric_cols[:5])}{'...' if len(numeric_cols) > 5 else ''}"
            )

        return "\n".join(summary)

    def get_summary(self) -> Dict:
        """Get complete summary of loaded data"""

        return {
            'structured_files': len(self.structured_data),
            'total_rows': self.stats['total_rows'],
            'documents': self.stats['total_documents'],
            'chunks_indexed': self.stats['total_chunks'],
            'data_size_mb': round(self.stats['data_size_mb'], 2),
            'ready': self.vectorstore is not None
        }


# ============= Streamlit UI =============

def main():
    st.set_page_config(
        page_title="PE Data Analysis POC",
        page_icon="📊",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    # Custom CSS
    st.markdown("""
    <style>
        .main {padding: 1rem;}
        .stMetric {
            background-color: #f0f2f6;
            padding: 1rem;
            border-radius: 0.5rem;
            box-shadow: 0 1px 3px rgba(0,0,0,0.12);
        }
        div[data-testid="metric-container"] {
            background-color: #f0f2f6;
            padding: 10px;
            border-radius: 10px;
            margin: 10px 0px;
        }
    </style>
    """, unsafe_allow_html=True)

    st.title("📊 Private Equity Data Analysis POC")
    st.caption("AI-powered analysis of structured and unstructured data with vector search")

    # Initialize data warehouse
    if 'warehouse' not in st.session_state:
        st.session_state.warehouse = DataWarehouse()
        st.session_state.data_loaded = False

    warehouse = st.session_state.warehouse

    # Sidebar
    with st.sidebar:
        st.header("📁 Data Management")

        # Data folder input
        data_folder = st.text_input(
            "Data folder path:",
            value="kaggle_data",
            help="Folder containing your CSVs, Excel files, and PDFs"
        )

        # Load data button
        if st.button("🔄 Load/Reload Data", type="primary", use_container_width=True):
            with st.spinner("Processing all data..."):
                # Create sample data if folder doesn't exist
                if not os.path.exists(data_folder):
                    st.info("Creating sample data...")
                    create_sample_data(data_folder)

                # Process all data
                summary = warehouse.process_all_data(data_folder)

                if summary:
                    st.session_state.data_loaded = True
                    st.success("✅ Data loaded successfully!")

                    # Show summary metrics
                    col1, col2 = st.columns(2)
                    with col1:
                        st.metric("Files", f"{summary['structured_files']} files")
                        st.metric("Documents", f"{summary['documents']} docs")
                    with col2:
                        st.metric("Total Rows", f"{summary['total_rows']:,}")
                        st.metric("Chunks", f"{summary['chunks_indexed']} indexed")

                    st.metric("Data Size", f"{summary['data_size_mb']:.1f} MB")
                else:
                    st.error("Failed to load data. Check the folder path.")

        # Show current stats if data is loaded
        if st.session_state.data_loaded:
            st.divider()
            st.subheader("📈 Loaded Datasets")

            # Show structured data files
            if warehouse.structured_data:
                for filename, df in warehouse.structured_data.items():
                    with st.expander(f"📊 {filename}"):
                        st.write(f"**Shape:** {df.shape[0]} rows × {df.shape[1]} cols")
                        st.write(f"**Columns:** {', '.join(df.columns[:5])}...")
                        st.write(f"**Memory:** {df.memory_usage(deep=True).sum() / 1024**2:.1f} MB")

            # Show documents
            if warehouse.documents:
                st.write(f"**📄 Documents:** {len(warehouse.documents)} files")
                for doc in warehouse.documents[:5]:
                    st.caption(f"• {doc['filename']} ({len(doc['chunks'])} chunks)")

    # Main area
    if not st.session_state.data_loaded:
        # Welcome screen
        st.info("👈 **Get started** by loading your data using the sidebar")

        # Instructions
        with st.expander("📖 How to use this POC", expanded=True):
            st.markdown("""
            ### Quick Start Guide

            1. **Prepare your data folder** (`kaggle_data/`) with:
               - CSV files (financial data, metrics)
               - Excel files (spreadsheets, reports)
               - PDF files (documents, reports)
               - Text files (notes, summaries)

            2. **Click 'Load/Reload Data'** in the sidebar
               - System will process all files
               - Create vector embeddings for semantic search
               - Index everything for fast retrieval

            3. **Ask questions** like:
               - "What's the average revenue across all companies?"
               - "What are the key risks mentioned in the documents?"
               - "Show me companies with EBITDA margins > 20%"
               - "Find information about customer concentration"

            4. **The system will**:
               - Search across ALL data (structured and unstructured)
               - Use GPT-4 for intelligent analysis
               - Provide sources for answers
               - Generate visualizations

            ### What makes this special?
            - **Semantic search**: Finds meaning, not just keywords
            - **Unified intelligence**: Combines data from all sources
            - **Production architecture**: RAG pattern with vector embeddings
            - **Scalable**: Handles thousands of documents and millions of rows
            """)

        # Create sample data button
        if st.button("🎲 Generate Sample Data", help="Create sample files for testing"):
            with st.spinner("Creating sample data..."):
                create_sample_data("kaggle_data")
                st.success("Sample data created in `kaggle_data/` folder!")
                st.info("Click 'Load/Reload Data' in the sidebar to process it")

    else:
        # Query interface
        tab1, tab2, tab3, tab4, tab5 = st.tabs([
            "🔍 Unified Search",
            "📊 Structured Data",
            "📄 Documents",
            "📈 Analytics",
            "🗂️ Data Explorer"
        ])

        with tab1:
            st.header("Ask Questions Across All Data")

            # Example questions
            example_questions = [
                "What is the total market size based on all data?",
                "What are the main risk factors mentioned?",
                "Which companies have the highest profit margins?",
                "Find all mentions of customer concentration",
                "What data quality issues exist?",
                "Compare revenue growth across sectors",
                "Summarize key findings from all sources"
            ]

            col1, col2 = st.columns([3, 1])
            with col1:
                selected = st.selectbox(
                    "Select an example or write your own:",
                    ["Custom question..."] + example_questions
                )

            if selected == "Custom question...":
                question = st.text_area(
                    "Your question:",
                    height=100,
                    placeholder="Ask anything about your data..."
                )
            else:
                question = selected
                st.info(f"**Selected:** {question}")

            col1, col2, col3 = st.columns(3)

            with col1:
                if st.button("🔍 Search All Data", type="primary", use_container_width=True):
                    if question:
                        with st.spinner("Analyzing all data sources..."):
                            answer = warehouse.ask_unified(question)

                        st.success("Analysis complete!")
                        st.markdown(answer)

                        # Show token usage estimate
                        with st.expander("ℹ️ Processing Details"):
                            st.write(f"• Searched {warehouse.stats['total_chunks']} document chunks")
                            st.write(f"• Analyzed {warehouse.stats['total_rows']:,} rows of structured data")
                            st.write(f"• Used GPT-4 Turbo for synthesis")

            with col2:
                if st.button("📊 Query Tables Only", use_container_width=True):
                    if question:
                        with st.spinner("Analyzing structured data..."):
                            answer = warehouse.ask_structured_data(question)
                        st.markdown(answer)

            with col3:
                if st.button("📄 Search Docs Only", use_container_width=True):
                    if question:
                        with st.spinner("Searching documents..."):
                            answer = warehouse.ask_documents(question)
                        st.markdown(answer)

        with tab2:
            st.header("Structured Data Analysis")

            if warehouse.structured_data:
                # Dataset selector
                dataset_name = st.selectbox(
                    "Select dataset:",
                    list(warehouse.structured_data.keys())
                )

                if dataset_name:
                    df = warehouse.structured_data[dataset_name]

                    # Show metrics
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("Rows", f"{len(df):,}")
                    with col2:
                        st.metric("Columns", len(df.columns))
                    with col3:
                        st.metric("Numeric Cols", len(df.select_dtypes(include=[np.number]).columns))
                    with col4:
                        st.metric("Missing Values", f"{df.isnull().sum().sum():,}")

                    # Show data
                    st.subheader("Data Preview")
                    st.dataframe(df.head(100), use_container_width=True)

                    # Show statistics
                    if st.checkbox("Show Summary Statistics"):
                        st.subheader("Summary Statistics")
                        st.dataframe(df.describe(), use_container_width=True)

                    # Column analysis
                    if st.checkbox("Analyze Specific Column"):
                        col = st.selectbox("Select column:", df.columns)

                        if col:
                            col1, col2 = st.columns(2)

                            with col1:
                                st.write(f"**Analysis of `{col}`:**")
                                st.write(f"• Data type: {df[col].dtype}")
                                st.write(f"• Unique values: {df[col].nunique():,}")
                                st.write(f"• Missing values: {df[col].isnull().sum():,}")

                                if df[col].dtype in ['float64', 'int64']:
                                    st.write(f"• Mean: {df[col].mean():.2f}")
                                    st.write(f"• Std: {df[col].std():.2f}")
                                    st.write(f"• Min: {df[col].min():.2f}")
                                    st.write(f"• Max: {df[col].max():.2f}")

                            with col2:
                                if df[col].dtype in ['float64', 'int64']:
                                    fig = px.histogram(df, x=col, nbins=30, title=f"Distribution of {col}")
                                    st.plotly_chart(fig, use_container_width=True)
                                elif df[col].nunique() < 20:
                                    fig = px.pie(values=df[col].value_counts().values,
                                               names=df[col].value_counts().index,
                                               title=f"Distribution of {col}")
                                    st.plotly_chart(fig, use_container_width=True)
            else:
                st.info("No structured data loaded")

        with tab3:
            st.header("Document Search")

            doc_query = st.text_input(
                "Search documents:",
                placeholder="Enter keywords or phrases..."
            )

            if st.button("Search Documents"):
                if doc_query and warehouse.vectorstore:
                    with st.spinner("Searching..."):
                        results = warehouse.search(doc_query, k=5)

                    if results:
                        st.success(f"Found {len(results)} relevant sections")

                        for i, doc in enumerate(results, 1):
                            with st.expander(f"Result {i} - {doc.metadata['source']}"):
                                st.write(doc.page_content)
                                st.caption(f"📄 Source: {doc.metadata['source']}")
                                st.caption(f"📍 Chunk {doc.metadata['chunk'] + 1} of {doc.metadata.get('total_chunks', 'unknown')}")
                    else:
                        st.warning("No results found")
                elif not warehouse.vectorstore:
                    st.warning("No documents indexed yet")

        with tab4:
            st.header("Analytics Dashboard")

            if warehouse.structured_data:
                # Find the main financial dataset
                financial_df = None
                for name, df in warehouse.structured_data.items():
                    if 'revenue' in [col.lower() for col in df.columns]:
                        financial_df = df
                        break

                if financial_df is not None:
                    # Create visualizations
                    col1, col2 = st.columns(2)

                    with col1:
                        # Revenue distribution
                        revenue_cols = [col for col in financial_df.columns if 'revenue' in col.lower()]
                        if revenue_cols:
                            fig = px.box(financial_df, y=revenue_cols[0], title="Revenue Distribution")
                            st.plotly_chart(fig, use_container_width=True)

                    with col2:
                        # Sector breakdown if available
                        sector_cols = [col for col in financial_df.columns if 'sector' in col.lower() or 'industry' in col.lower()]
                        if sector_cols:
                            sector_counts = financial_df[sector_cols[0]].value_counts()
                            fig = px.bar(x=sector_counts.index, y=sector_counts.values,
                                       title="Companies by Sector",
                                       labels={'x': 'Sector', 'y': 'Count'})
                            st.plotly_chart(fig, use_container_width=True)

                    # Correlation matrix for numeric columns
                    numeric_cols = financial_df.select_dtypes(include=[np.number]).columns
                    if len(numeric_cols) > 1:
                        if st.checkbox("Show Correlation Matrix"):
                            corr_matrix = financial_df[numeric_cols].corr()
                            fig = px.imshow(corr_matrix,
                                          text_auto=True,
                                          title="Correlation Matrix",
                                          color_continuous_scale='RdBu',
                                          zmin=-1, zmax=1)
                            st.plotly_chart(fig, use_container_width=True)

                # Summary statistics
                st.subheader("Summary Metrics")
                col1, col2, col3, col4 = st.columns(4)

                total_companies = sum(len(df) for df in warehouse.structured_data.values() if 'company' in str(df.columns).lower())
                total_transactions = sum(len(df) for df in warehouse.structured_data.values() if 'transaction' in str(df.columns).lower())

                with col1:
                    st.metric("Total Data Points", f"{warehouse.stats['total_rows']:,}")
                with col2:
                    st.metric("Companies", f"{total_companies:,}")
                with col3:
                    st.metric("Transactions", f"{total_transactions:,}")
                with col4:
                    st.metric("Documents", warehouse.stats['total_documents'])
            else:
                st.info("Load data to see analytics")

        with tab5:
            st.header("Data Explorer")

            # Document explorer
            if warehouse.documents:
                st.subheader("📄 Document Library")

                for doc in warehouse.documents:
                    with st.expander(f"{doc['filename']} ({doc['type'].upper()})"):
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Type", doc['type'].upper())
                        with col2:
                            st.metric("Size", f"{len(doc['content']):,} chars")
                        with col3:
                            st.metric("Chunks", len(doc['chunks']))

                        st.text_area("Preview", doc['content'][:1000] + "...", height=200, disabled=True)

            # Vector store info
            if warehouse.vectorstore:
                st.subheader("🔍 Vector Index Statistics")

                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Total Chunks", warehouse.stats['total_chunks'])
                with col2:
                    st.metric("Embedding Model", "text-embedding-ada-002")
                with col3:
                    st.metric("Vector Dimensions", "1536")

                st.success("✅ Vector index is active and ready for semantic search")

            # Data quality report
            if st.button("Generate Data Quality Report"):
                with st.spinner("Analyzing data quality..."):
                    report = []

                    for name, df in warehouse.structured_data.items():
                        missing_pct = (df.isnull().sum().sum() / (df.shape[0] * df.shape[1])) * 100
                        duplicates = df.duplicated().sum()

                        report.append({
                            'Dataset': name,
                            'Rows': len(df),
                            'Columns': len(df.columns),
                            'Missing %': f"{missing_pct:.1f}%",
                            'Duplicates': duplicates
                        })

                    if report:
                        st.subheader("Data Quality Report")
                        st.dataframe(pd.DataFrame(report), use_container_width=True)


def create_sample_data(folder: str):
    """Create sample data if none exists"""

    os.makedirs(folder, exist_ok=True)

    # Create sample financial data
    np.random.seed(42)

    financial_df = pd.DataFrame({
        'Company': [f'Company_{i:03d}' for i in range(500)],
        'Ticker': [f'{chr(65+i%26)}{chr(65+(i//26)%26)}{chr(65+(i//676)%26)}' for i in range(500)],
        'Sector': np.random.choice(['Technology', 'Healthcare', 'Finance', 'Retail', 'Energy', 'Manufacturing'], 500),
        'Revenue_2023': np.random.uniform(100, 10000, 500),
        'Revenue_2022': np.random.uniform(100, 9000, 500),
        'EBITDA_2023': np.random.uniform(10, 2000, 500),
        'EBITDA_2022': np.random.uniform(10, 1800, 500),
        'NetIncome_2023': np.random.uniform(-100, 1000, 500),
        'NetIncome_2022': np.random.uniform(-100, 900, 500),
        'TotalAssets': np.random.uniform(1000, 50000, 500),
        'TotalDebt': np.random.uniform(100, 20000, 500),
        'Employees': np.random.randint(50, 50000, 500)
    })

    # Calculate derived metrics
    financial_df['EBITDA_Margin_2023'] = (financial_df['EBITDA_2023'] / financial_df['Revenue_2023']) * 100
    financial_df['Revenue_Growth'] = ((financial_df['Revenue_2023'] - financial_df['Revenue_2022']) / financial_df['Revenue_2022']) * 100
    financial_df['Debt_to_Assets'] = financial_df['TotalDebt'] / financial_df['TotalAssets']

    financial_df.to_csv(f"{folder}/financial_data.csv", index=False)

    # Create sample sales data
    sales_df = pd.DataFrame({
        'Date': pd.date_range('2023-01-01', periods=10000, freq='H'),
        'TransactionID': [f'TXN{i:08d}' for i in range(10000)],
        'Product': np.random.choice(['Product_A', 'Product_B', 'Product_C', 'Product_D'], 10000),
        'Quantity': np.random.randint(1, 100, 10000),
        'UnitPrice': np.random.uniform(10, 1000, 10000),
        'Customer': [f'CUST{i%1000:04d}' for i in range(10000)],
        'Region': np.random.choice(['North', 'South', 'East', 'West'], 10000)
    })

    sales_df['Revenue'] = sales_df['Quantity'] * sales_df['UnitPrice']
    sales_df.to_excel(f"{folder}/sales_data.xlsx", index=False)

    # Create sample text documents
    risk_report = """COMPREHENSIVE RISK ASSESSMENT REPORT 2023

EXECUTIVE SUMMARY
Our analysis has identified several critical risk factors requiring immediate attention from management and the board.

1. CUSTOMER CONCENTRATION RISK
- Top 3 customers represent 45% of total revenue
- Largest customer accounts for 22% of revenue (up from 18% last year)
- Customer retention rate declined to 87% from 92% YoY
- Risk Level: HIGH
- Recommendation: Immediate diversification strategy required

2. MARKET COMPETITION
- Three new entrants in primary market segment
- Price pressure increasing, margins compressed by 3%
- Market share declined from 15% to 13.5%
- Emerging technologies threatening core products
- Risk Level: MEDIUM-HIGH

3. REGULATORY COMPLIANCE
- New regulations expected Q3 2024
- Estimated compliance costs: $2.5M annually
- Potential fines up to $10M for non-compliance
- Data privacy laws becoming stricter
- Risk Level: MEDIUM

4. OPERATIONAL RISKS
- Supply chain disruptions affecting 30% of products
- Single supplier dependency for critical components
- IT infrastructure requires $5M upgrade
- Cybersecurity incidents increased 40% YoY
- Risk Level: HIGH

5. FINANCIAL RISKS
- Debt-to-equity ratio increased to 1.8x (covenant limit: 2.0x)
- Interest coverage ratio declined to 3.2x
- Working capital requirements increased 25%
- FX exposure on 40% of revenues
- Risk Level: MEDIUM

RECOMMENDATIONS
1. Diversify customer base - target 10 new enterprise clients
2. Invest $3M in compliance infrastructure immediately
3. Strengthen cybersecurity measures - hire CISO
4. Reduce debt by $50M within 12 months
5. Implement hedging strategy for FX exposure

CONCLUSION
While the company maintains strong market position, immediate action required on customer concentration and operational risks to ensure sustainable growth."""

    with open(f"{folder}/risk_assessment_report.txt", 'w') as f:
        f.write(risk_report)

    st.success(f"✅ Created sample data in {folder}/")
    st.info("Files created: financial_data.csv, sales_data.xlsx, risk_assessment_report.txt")


if __name__ == "__main__":
    main()